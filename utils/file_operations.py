"""
file_operations.py - Handles file import/export and report generation
"""
from datetime import datetime
import csv
import tempfile
import os
import webbrowser

# Try to import pandas, but provide alternatives if not available
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("Warning: pandas module not found. Using basic CSV handling instead.")

class FileOperations:
    """
    Handles file operations for the test prioritization application
    """
    @staticmethod
    def export_to_csv(filename, tests, factors):
        """
        Export tests to a CSV file
        
        Args:
            filename (str): Path to the CSV file
            tests (list): List of test dictionaries to export
            factors (dict): Dictionary of factors
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Sort tests by score (descending)
            sorted_tests = sorted(tests, key=lambda x: x.get("total_score", 0), reverse=True)
            
            # Create field names (column headers)
            field_names = ["Rank", "Priority", "Ticket ID", "Section", "Test Name", "Description", 
                            "Total Score (100-point)", "Raw Score", "Test ID"]
            
            # Add factor score headers
            for factor_key, factor_info in factors.items():
                field_names.append(factor_info["name"])
            
            # Check if we have yes/no questions to add to the CSV
            has_yes_no = False
            if len(sorted_tests) > 0 and 'yes_no_answers' in sorted_tests[0]:
                has_yes_no = True
                # Get the question texts from the first test's yes_no_answers
                for question_key in sorted_tests[0].get('yes_no_answers', {}):
                    # Fallback to just using the key as header
                    field_names.append(f"Question: {question_key}")
            
            # Prepare data for export
            data = []
            for i, test in enumerate(sorted_tests):
                # Ensure description is never None or NaN
                description = test.get("description", "")
                if description is None or description == "nan" or (hasattr(description, "lower") and description.lower() == "nan"):
                    description = ""
                
                # Create row dictionary    
                row = {
                    "Rank": i + 1,
                    "Priority": test.get("priority", ""),
                    "Ticket ID": test.get("ticket_id", "N/A"),
                    "Section": test.get("section", ""),
                    "Test Name": test.get("name", ""),
                    "Description": description,
                    "Total Score (100-point)": test.get("total_score", 0),
                    "Raw Score": test.get("raw_score", "N/A")
                    
                }
                
                # Add factor scores
                for factor in factors:
                    factor_name = factors[factor]["name"]
                    # Get scores safely, defaulting to 0 if not present
                    scores = test.get("scores", {})
                    row[factor_name] = scores.get(factor, 0)
                
                # Add yes/no answers if available
                if has_yes_no and 'yes_no_answers' in test:
                    for question_key, answer in test['yes_no_answers'].items():
                        question_text = f"Question: {question_key}"
                        row[question_text] = "Yes" if answer else "No"
                
                row["Test ID"] = test.get("id", "")

                data.append(row)
            
            # Use pandas if available, otherwise use the built-in csv module
            if PANDAS_AVAILABLE:
                # Create DataFrame and export
                df = pd.DataFrame(data)
                df.to_csv(filename, index=False)
            else:
                # Use csv module
                with open(filename, 'w', newline='') as csvfile:
                    writer = csv.DictWriter(csvfile, fieldnames=field_names)
                    writer.writeheader()
                    for row in data:
                        writer.writerow(row)
            
            return True
        
        except Exception as e:
            print(f"Export error: {str(e)}")
            return False
    
    @staticmethod
    def import_from_csv(filename):
        """
        Import tests from a CSV file
        
        Args:
            filename (str): Path to the CSV file
            
        Returns:
            tuple: (success, data or error message)
                success (bool): True if successful, False otherwise
                data (list): List of dictionaries if successful, error message otherwise
        """
        try:
            if PANDAS_AVAILABLE:
                # Use pandas if available
                df = pd.read_csv(filename)
                data = df.to_dict(orient='records')
            else:
                # Use csv module if pandas not available
                data = []
                with open(filename, 'r', newline='') as csvfile:
                    reader = csv.DictReader(csvfile)
                    for row in reader:
                        data.append(dict(row))
            
            return True, data
        
        except Exception as e:
            error_message = f"Import error: {str(e)}"
            print(error_message)
            return False, error_message
    
    @staticmethod
    def generate_report_text(tests, priority_tiers, model=None):
        """
        Generate text for the prioritization report
        
        Args:
            tests (list): List of all test dictionaries
            priority_tiers (dict): Dictionary with priority tier tests
            model: The prioritization model (added parameter)
            
        Returns:
            str: Formatted report text
        """
        # Report header
        header = f"TEST AUTOMATION PRIORITY REPORT\n"
        header += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        header += f"Total Tests: {len(tests)}\n"
        
        # Group tests by section
        sections = {}
        for test in tests:
            section = test.get("section", "")
            if section not in sections:
                sections[section] = []
            sections[section].append(test)
        
        if sections:
            header += f"Sections: {len(sections)}\n"
        
        header += "=" * 70 + "\n\n"
        
        report_text = header
        
        # Get the tiers
        highest_priority = priority_tiers["highest"]
        high_priority = priority_tiers["high"]
        medium_priority = priority_tiers["medium"]
        low_priority = priority_tiers["low"]
        lowest_priority = priority_tiers["lowest"]
        cant_automate = priority_tiers.get("cant_automate", [])  # Get "Can't Automate" tests if available

        highest_threshold = priority_tiers["highest_threshold"]
        high_threshold = priority_tiers["high_threshold"]
        medium_threshold = priority_tiers["medium_threshold"]
        low_threshold = priority_tiers["low_threshold"]
        lowest_threshold = priority_tiers["lowest_threshold"]

        # Highest priority section
        report_text += f"HIGHEST PRIORITY TESTS (Score >= {highest_threshold:.1f}):\n"
        report_text += f"Recommended for immediate automation\n"
        report_text += "-" * 70 + "\n"
        
        for i, test in enumerate(highest_priority):
            report_text += f"| {i+1}. {test['name']} (ID: {test['id']})\n"
            report_text += f"|    Score: {test['total_score']:.1f}\n"

            # Add section if available
            if test.get("section"):
                report_text += f"|    Section: {test['section']}\n"

            # Add description if available
            if 'description' in test:
                report_text += f"|    Description: {test['description']}\n"
            
            # Add score details with descriptions
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                report_text += f"|    Factor Scores:\n"
                for factor, score in test['scores'].items():
                    if factor in model.factors and score in model.score_options.get(factor, {}):
                        factor_name = model.factors[factor]["name"]
                        score_description = model.score_options[factor][score]
                        report_text += f"|      - {factor_name}: {score} - {score_description}\n"
            
            # Add yes/no answers if available
            if 'yes_no_answers' in test:
                for key, answer in test['yes_no_answers'].items():
                    report_text += f"|    * {key}: {answer}\n"
            
            report_text += "|\n"

        report_text += "-" * 70 + "\n"
        
        report_text += "\n"
        
        # High priority section
        report_text += f"HIGH PRIORITY TESTS (Score {high_threshold:.1f} - {highest_threshold:.1f}):\n"
        report_text += f"Recommended for second phase automation\n"
        report_text += "-" * 70 + "\n"
        
        for i, test in enumerate(high_priority):
            report_text += f"| {i+1}. {test['name']} (ID: {test['id']})\n"
            report_text += f"|    Score: {test['total_score']:.1f}\n"

            # Add section if available
            if test.get("section"):
                report_text += f"|    Section: {test['section']}\n"

            # Add description if available
            if 'description' in test:
                report_text += f"|    Description: {test['description']}\n"
            
            # Add score details with descriptions
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                report_text += f"|    Factor Scores:\n"
                for factor, score in test['scores'].items():
                    if factor in model.factors and score in model.score_options.get(factor, {}):
                        factor_name = model.factors[factor]["name"]
                        score_description = model.score_options[factor][score]
                        report_text += f"|      - {factor_name}: {score} - {score_description}\n"
            
            # Add yes/no answers if available
            if 'yes_no_answers' in test:
                for key, answer in test['yes_no_answers'].items():
                    report_text += f"|    * {key}: {answer}\n"
            
            report_text += "|\n"

        report_text += "-" * 70 + "\n"
        
        report_text += "\n"
        
        # Medium priority section
        report_text += f"MEDIUM PRIORITY TESTS (Score {medium_threshold:.1f} - {high_threshold:.1f}):\n"
        report_text += f"Recommended for third phase automation\n"
        report_text += "-" * 70 + "\n"
        
        for i, test in enumerate(medium_priority):
            report_text += f"| {i+1}. {test['name']} (ID: {test['id']})\n"
            report_text += f"|    Score: {test['total_score']:.1f}\n"

            # Add section if available
            if test.get("section"):
                report_text += f"|    Section: {test['section']}\n"

            # Add description if available
            if 'description' in test:
                report_text += f"|    Description: {test['description']}\n"
            
            # Add score details with descriptions
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                report_text += f"|    Factor Scores:\n"
                for factor, score in test['scores'].items():
                    if factor in model.factors and score in model.score_options.get(factor, {}):
                        factor_name = model.factors[factor]["name"]
                        score_description = model.score_options[factor][score]
                        report_text += f"|      - {factor_name}: {score} - {score_description}\n"
            
            # Add yes/no answers if available
            if 'yes_no_answers' in test:
                for key, answer in test['yes_no_answers'].items():
                    report_text += f"|    * {key}: {answer}\n"
            
            report_text += "|\n"

        report_text += "-" * 70 + "\n"
        
        report_text += "\n"
        
        # Low priority section
        report_text += f"LOW PRIORITY TESTS (Score {low_threshold:.1f} - {medium_threshold:.1f}):\n"
        report_text += f"Consider for later phases or keep as manual tests\n"
        report_text += "-" * 70 + "\n"
        
        for i, test in enumerate(low_priority):
            report_text += f"| {i+1}. {test['name']} (ID: {test['id']})\n"
            report_text += f"|    Score: {test['total_score']:.1f}\n"

            # Add section if available
            if test.get("section"):
                report_text += f"|    Section: {test['section']}\n"

            # Add description if available
            if 'description' in test:
                report_text += f"|    Description: {test['description']}\n"

            # Add score details with descriptions
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                report_text += f"|    Factor Scores:\n"
                for factor, score in test['scores'].items():
                    if factor in model.factors and score in model.score_options.get(factor, {}):
                        factor_name = model.factors[factor]["name"]
                        score_description = model.score_options[factor][score]
                        report_text += f"|      - {factor_name}: {score} - {score_description}\n"

            if 'yes_no_answers' in test:
                for key, answer in test['yes_no_answers'].items():
                    report_text += f"|    * {key}: {answer}\n"
                    
            report_text += "|\n"

        report_text += "-" * 70 + "\n"

        report_text += "\n"

        # Lowest priority section
        report_text += f"LOWEST PRIORITY TESTS (Score <= {low_threshold:.1f}):\n"
        report_text += f"Not Recommended for automation\n"
        report_text += "-" * 70 + "\n"
        
        for i, test in enumerate(lowest_priority):
            report_text += f"| {i+1}. {test['name']} (ID: {test['id']})\n"
            report_text += f"|    Score: {test['total_score']:.1f}\n"
            
            # Add section if available
            if test.get("section"):
                report_text += f"|    Section: {test['section']}\n"
            
            # Add description if available
            if 'description' in test:
                report_text += f"|    Description: {test['description']}\n"
            
            # Add score details with descriptions
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                report_text += f"|    Factor Scores:\n"
                for factor, score in test['scores'].items():
                    if factor in model.factors and score in model.score_options.get(factor, {}):
                        factor_name = model.factors[factor]["name"]
                        score_description = model.score_options[factor][score]
                        report_text += f"|      - {factor_name}: {score} - {score_description}\n"
            
            # Add yes/no answers if available
            if 'yes_no_answers' in test:
                for key, answer in test['yes_no_answers'].items():
                    report_text += f"|    * {key}: {answer}\n"
            
            report_text += "|\n"
        
        report_text += "-" * 70 + "\n"
        
        report_text += "\n"
        
        # Can't Automate section
        if cant_automate:
            report_text += f"TESTS THAT CAN'T BE AUTOMATED:\n"
            report_text += f"These tests have been identified as not possible to automate\n"
            report_text += "-" * 70 + "\n"
            
            for i, test in enumerate(cant_automate):
                report_text += f"| {i+1}. {test['name']} (ID: {test['id']})\n"
                
                # Add section if available
                if test.get("section"):
                    report_text += f"|    Section: {test['section']}\n"
                
                # Add description if available
                if 'description' in test:
                    report_text += f"|    Description: {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"|    Factor Scores:\n"
                    # First show the "Can it be automated?" factor to explain why it's in this category
                    if "can_be_automated" in test['scores'] and test['scores']["can_be_automated"] == 1:
                        factor_name = model.factors["can_be_automated"]["name"]
                        score_description = model.score_options["can_be_automated"][1]
                        report_text += f"|      - {factor_name}: 1 - {score_description}\n"
                        
                    # Then show other factors
                    for factor, score in test['scores'].items():
                        if factor != "can_be_automated" and factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"|      - {factor_name}: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if 'yes_no_answers' in test:
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"|    * {key}: {answer}\n"
                
                report_text += "|\n"
            
            report_text += "-" * 70 + "\n"
            
        # Add section breakdown report
        if len(sections) > 1:  # Only add section breakdown if there's more than one section
            report_text += "\n"
            report_text += "SECTION BREAKDOWN:\n"
            report_text += "-" * 70 + "\n"
            
            for section_name, section_tests in sorted(sections.items()):
                # Skip empty section name
                if not section_name:
                    continue
                    
                # Count tests by priority in this section
                priority_counts = {"Highest": 0, "High": 0, "Medium": 0, "Low": 0, "Lowest": 0, "Can't Automate": 0}
                for test in section_tests:
                    priority = test.get("priority", "")
                    if priority in priority_counts:
                        priority_counts[priority] += 1
                
                report_text += f"Section: {section_name}\n"
                report_text += f"Total Tests: {len(section_tests)}\n"
                report_text += "Priority Distribution:\n"
                for priority, count in priority_counts.items():
                    if count > 0:
                        report_text += f"  - {priority}: {count} tests\n"
                report_text += "\n"
            
            report_text += "-" * 70 + "\n"
        
        return report_text
    
    @staticmethod
    def generate_markdown_report(tests, priority_tiers, model=None):
        """
        Generate markdown text for the prioritization report
        
        Args:
            tests (list): List of all test dictionaries
            priority_tiers (dict): Dictionary with priority tier tests
            model: The prioritization model
            
        Returns:
            str: Formatted markdown report text
        """
        from datetime import datetime
        
        # Report header
        header = f"# TEST AUTOMATION PRIORITY REPORT\n\n"
        header += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        header += f"**Total Tests:** {len(tests)}\n"
        
        # Group tests by section
        sections = {}
        for test in tests:
            section = test.get("section", "")
            if section not in sections:
                sections[section] = []
            sections[section].append(test)
        
        if sections:
            header += f"**Sections:** {len(sections)}\n"
        
        header += "\n---\n\n"
        
        report_text = header
        
        # Get the tiers
        highest_priority = priority_tiers["highest"]
        high_priority = priority_tiers["high"]
        medium_priority = priority_tiers["medium"]
        low_priority = priority_tiers["low"]
        lowest_priority = priority_tiers["lowest"]
        cant_automate = priority_tiers.get("cant_automate", [])

        highest_threshold = priority_tiers["highest_threshold"]
        high_threshold = priority_tiers["high_threshold"]
        medium_threshold = priority_tiers["medium_threshold"]
        low_threshold = priority_tiers["low_threshold"]
        lowest_threshold = priority_tiers["lowest_threshold"]

        # Highest priority section
        report_text += f"## 🔴 HIGHEST PRIORITY TESTS (Score >= {highest_threshold:.1f})\n\n"
        report_text += f"*Recommended for immediate automation*\n\n"
        
        if highest_priority:
            for i, test in enumerate(highest_priority):
                report_text += f"### {i+1}. {test['name']}\n"
                report_text += f"**Score:** {test['total_score']:.1f}\n"
                report_text += f"**Ticket:** {test['ticket_id']}\n"

                # Add section if available
                if test.get("section"):
                    report_text += f"**Section:** {test['section']}\n"

                # Add description if available
                if test.get('description'):
                    report_text += f"**Description:** {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"**Factor Scores:**\n"
                    for factor, score in test['scores'].items():
                        if factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"* **{factor_name}**: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if test.get('yes_no_answers'):
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"* **{key}**: {answer}\n"
                
                report_text += "\n"
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # High priority section
        report_text += f"## 🟠 HIGH PRIORITY TESTS (Score {high_threshold:.1f} - {highest_threshold:.1f})\n\n"
        report_text += f"*Recommended for second phase automation*\n\n"
        
        if high_priority:
            for i, test in enumerate(high_priority):
                report_text += f"### {i+1}. {test['name']}\n"
                report_text += f"**Score:** {test['total_score']:.1f}\n"
                report_text += f"**Ticket:** {test['ticket_id']}\n"

                # Add section if available
                if test.get("section"):
                    report_text += f"**Section:** {test['section']}\n"

                # Add description if available
                if test.get('description'):
                    report_text += f"**Description:** {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"**Factor Scores:**\n"
                    for factor, score in test['scores'].items():
                        if factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"* **{factor_name}**: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if test.get('yes_no_answers'):
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"* **{key}**: {answer}\n"
                
                report_text += "\n"
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Medium priority section
        report_text += f"## 🟡 MEDIUM PRIORITY TESTS (Score {medium_threshold:.1f} - {high_threshold:.1f})\n\n"
        report_text += f"*Recommended for third phase automation*\n\n"
        
        if medium_priority:
            for i, test in enumerate(medium_priority):
                report_text += f"### {i+1}. {test['name']}\n"
                report_text += f"**Score:** {test['total_score']:.1f}\n"
                report_text += f"**Ticket:** {test['ticket_id']}\n"

                # Add section if available
                if test.get("section"):
                    report_text += f"**Section:** {test['section']}\n"

                # Add description if available
                if test.get('description'):
                    report_text += f"**Description:** {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"**Factor Scores:**\n"
                    for factor, score in test['scores'].items():
                        if factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"* **{factor_name}**: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if test.get('yes_no_answers'):
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"* **{key}**: {answer}\n"
                
                report_text += "\n"
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Low priority section
        report_text += f"## 🔵 LOW PRIORITY TESTS (Score {low_threshold:.1f} - {medium_threshold:.1f})\n\n"
        report_text += f"*Consider for later phases or keep as manual tests*\n\n"
        
        if low_priority:
            for i, test in enumerate(low_priority):
                report_text += f"### {i+1}. {test['name']}\n"
                report_text += f"**Score:** {test['total_score']:.1f}\n"
                report_text += f"**Ticket:** {test['ticket_id']}\n"

                # Add section if available
                if test.get("section"):
                    report_text += f"**Section:** {test['section']}\n"

                # Add description if available
                if test.get('description'):
                    report_text += f"**Description:** {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"**Factor Scores:**\n"
                    for factor, score in test['scores'].items():
                        if factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"* **{factor_name}**: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if test.get('yes_no_answers'):
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"* **{key}**: {answer}\n"
                
                report_text += "\n"
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Lowest priority section
        report_text += f"## 🔷 LOWEST PRIORITY TESTS (Score <= {low_threshold:.1f})\n\n"
        report_text += f"*Not Recommended for automation*\n\n"
        
        if lowest_priority:
            for i, test in enumerate(lowest_priority):
                report_text += f"### {i+1}. {test['name']}\n"
                report_text += f"**Score:** {test['total_score']:.1f}\n"
                report_text += f"**Ticket:** {test['ticket_id']}\n"

                # Add section if available
                if test.get("section"):
                    report_text += f"**Section:** {test['section']}\n"

                # Add description if available
                if test.get('description'):
                    report_text += f"**Description:** {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"**Factor Scores:**\n"
                    for factor, score in test['scores'].items():
                        if factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"* **{factor_name}**: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if test.get('yes_no_answers'):
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"* **{key}**: {answer}\n"
                
                report_text += "\n"
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Can't Automate section
        if cant_automate:
            report_text += f"## ⚪ TESTS THAT CAN'T BE AUTOMATED\n\n"
            report_text += f"*These tests have been identified as not possible to automate*\n\n"
            
            for i, test in enumerate(cant_automate):
                report_text += f"### {i+1}. {test['name']}\n"
                report_text += f"**Ticket:** {test['ticket_id']}\n"
                
                # Add section if available
                if test.get("section"):
                    report_text += f"**Section:** {test['section']}\n"
                
                # Add description if available
                if test.get('description'):
                    report_text += f"**Description:** {test['description']}\n"
                
                # Add score details with descriptions
                if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                    report_text += f"**Factor Scores:**\n"
                    # First show the "Can it be automated?" factor to explain why it's in this category
                    if "can_be_automated" in test['scores'] and test['scores']["can_be_automated"] == 1:
                        factor_name = model.factors["can_be_automated"]["name"]
                        score_description = model.score_options["can_be_automated"][1]
                        report_text += f"* **{factor_name}**: 1 - {score_description}\n"
                        
                    # Then show other factors
                    for factor, score in test['scores'].items():
                        if factor != "can_be_automated" and factor in model.factors and score in model.score_options.get(factor, {}):
                            factor_name = model.factors[factor]["name"]
                            score_description = model.score_options[factor][score]
                            report_text += f"* **{factor_name}**: {score} - {score_description}\n"
                
                # Add yes/no answers if available
                if test.get('yes_no_answers'):
                    for key, answer in test['yes_no_answers'].items():
                        report_text += f"* **{key}**: {answer}\n"
                
                report_text += "\n"
            
            report_text += "---\n\n"
        
        # Add section breakdown report
        if len(sections) > 1:  # Only add section breakdown if there's more than one section
            report_text += "## SECTION BREAKDOWN\n\n"
            
            for section_name, section_tests in sorted(sections.items()):
                # Skip empty section name
                if not section_name:
                    continue
                    
                # Count tests by priority in this section
                priority_counts = {"Highest": 0, "High": 0, "Medium": 0, "Low": 0, "Lowest": 0, "Can't Automate": 0}
                for test in section_tests:
                    priority = test.get("priority", "")
                    if priority in priority_counts:
                        priority_counts[priority] += 1
                
                report_text += f"### Section: {section_name}\n"
                report_text += f"**Total Tests:** {len(section_tests)}\n"
                report_text += "**Priority Distribution:**\n"
                for priority, count in priority_counts.items():
                    if count > 0:
                        # Add emoji based on priority
                        emoji = ""
                        if priority == "Highest":
                            emoji = "🔴"
                        elif priority == "High":
                            emoji = "🟠"
                        elif priority == "Medium":
                            emoji = "🟡"
                        elif priority == "Low":
                            emoji = "🔵"
                        elif priority == "Lowest":
                            emoji = "🔷"
                        elif priority == "Can't Automate":
                            emoji = "⚪"
                        
                        report_text += f"* {emoji} **{priority}**: {count} tests\n"
                report_text += "\n"
            
            report_text += "---\n\n"
        
        return report_text
    
    @staticmethod
    def generate_enhanced_markdown_report(tests, priority_tiers, model=None):
        """
        Generate enhanced markdown text for the prioritization report with better HTML compatibility
        
        Args:
            tests (list): List of all test dictionaries
            priority_tiers (dict): Dictionary with priority tier tests
            model: The prioritization model
                
        Returns:
            str: Formatted markdown report text
        """
        from datetime import datetime
        
        # Report header
        header = f"# TEST AUTOMATION PRIORITY REPORT\n\n"
        header += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        header += f"**Total Tests:** {len(tests)}\n"
        
        # Group tests by section
        sections = {}
        for test in tests:
            section = test.get("section", "")
            if section not in sections:
                sections[section] = []
            sections[section].append(test)
        
        if sections:
            header += f"**Sections:** {len(sections)}\n"
        
        header += "\n---\n\n"
        
        # Add priority distribution summary
        highest_count = len(priority_tiers["highest"])
        high_count = len(priority_tiers["high"])
        medium_count = len(priority_tiers["medium"])
        low_count = len(priority_tiers["low"])
        lowest_count = len(priority_tiers["lowest"])
        cant_automate_count = len(priority_tiers.get("cant_automate", []))
        
        summary = "## Priority Distribution Summary\n\n"
        summary += "<div class='summary'>\n"
        summary += "<div class='summary-grid'>\n"
        summary += f"<div class='summary-item highest'><div class='summary-count'>{highest_count}</div><div class='summary-label'>Highest</div></div>\n"
        summary += f"<div class='summary-item high'><div class='summary-count'>{high_count}</div><div class='summary-label'>High</div></div>\n"
        summary += f"<div class='summary-item medium'><div class='summary-count'>{medium_count}</div><div class='summary-label'>Medium</div></div>\n"
        summary += f"<div class='summary-item low'><div class='summary-count'>{low_count}</div><div class='summary-label'>Low</div></div>\n"
        summary += f"<div class='summary-item lowest'><div class='summary-count'>{lowest_count}</div><div class='summary-label'>Lowest</div></div>\n"
        summary += f"<div class='summary-item cant-automate'><div class='summary-count'>{cant_automate_count}</div><div class='summary-label'>Can't Automate</div></div>\n"
        summary += "</div>\n</div>\n\n"
        
        report_text = header + summary
        
        # Get the tiers
        highest_priority = priority_tiers["highest"]
        high_priority = priority_tiers["high"]
        medium_priority = priority_tiers["medium"]
        low_priority = priority_tiers["low"]
        lowest_priority = priority_tiers["lowest"]
        cant_automate = priority_tiers.get("cant_automate", [])

        highest_threshold = priority_tiers["highest_threshold"]
        high_threshold = priority_tiers["high_threshold"]
        medium_threshold = priority_tiers["medium_threshold"]
        low_threshold = priority_tiers["low_threshold"]
        lowest_threshold = priority_tiers["lowest_threshold"]

        # Highest priority section
        report_text += f"## <div class='section-header section-highest'>🔴 HIGHEST PRIORITY TESTS (Score >= {highest_threshold:.1f})</div>\n\n"
        report_text += f"*Recommended for immediate automation*\n\n"
        
        if highest_priority:
            for i, test in enumerate(highest_priority):
                report_text += FileOperations.generate_test_card_markdown(test, i+1, "highest", model)
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # High priority section
        report_text += f"## <div class='section-header section-high'>🟠 HIGH PRIORITY TESTS (Score {high_threshold:.1f} - {highest_threshold:.1f})</div>\n\n"
        report_text += f"*Recommended for second phase automation*\n\n"
        
        if high_priority:
            for i, test in enumerate(high_priority):
                report_text += FileOperations.generate_test_card_markdown(test, i+1, "high", model)
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Medium priority section
        report_text += f"## <div class='section-header section-medium'>🟡 MEDIUM PRIORITY TESTS (Score {medium_threshold:.1f} - {high_threshold:.1f})</div>\n\n"
        report_text += f"*Recommended for third phase automation*\n\n"
        
        if medium_priority:
            for i, test in enumerate(medium_priority):
                report_text += FileOperations.generate_test_card_markdown(test, i+1, "medium", model)
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Low priority section
        report_text += f"## <div class='section-header section-low'>🔵 LOW PRIORITY TESTS (Score {low_threshold:.1f} - {medium_threshold:.1f})</div>\n\n"
        report_text += f"*Consider for later phases or keep as manual tests*\n\n"
        
        if low_priority:
            for i, test in enumerate(low_priority):
                report_text += FileOperations.generate_test_card_markdown(test, i+1, "low", model)
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Lowest priority section
        report_text += f"## <div class='section-header section-lowest'>🔷 LOWEST PRIORITY TESTS (Score <= {low_threshold:.1f})</div>\n\n"
        report_text += f"*Not Recommended for automation*\n\n"
        
        if lowest_priority:
            for i, test in enumerate(lowest_priority):
                report_text += FileOperations.generate_test_card_markdown(test, i+1, "lowest", model)
        else:
            report_text += "*No tests in this category*\n\n"
        
        report_text += "---\n\n"
        
        # Can't Automate section
        if cant_automate:
            report_text += f"## <div class='section-header section-cant-automate'>⚪ TESTS THAT CAN'T BE AUTOMATED</div>\n\n"
            report_text += f"*These tests have been identified as not possible to automate*\n\n"
            
            for i, test in enumerate(cant_automate):
                report_text += FileOperations.generate_test_card_markdown(test, i+1, "cant-automate", model)
            
            report_text += "---\n\n"
        
        # Add section breakdown report
        if len(sections) > 1:  # Only add section breakdown if there's more than one section
            report_text += "## SECTION BREAKDOWN\n\n"
            
            for section_name, section_tests in sorted(sections.items()):
                # Skip empty section name
                if not section_name:
                    continue
                    
                # Count tests by priority in this section
                priority_counts = {"Highest": 0, "High": 0, "Medium": 0, "Low": 0, "Lowest": 0, "Can't Automate": 0}
                for test in section_tests:
                    priority = test.get("priority", "")
                    if priority in priority_counts:
                        priority_counts[priority] += 1
                
                report_text += f"<div class='breakdown-item'>\n"
                report_text += f"<div class='breakdown-title'>Section: {section_name}</div>\n"
                report_text += f"<p><strong>Total Tests:</strong> {len(section_tests)}</p>\n"
                report_text += "<div class='breakdown-stats'>\n"
                
                # Add emoji based on priority
                for priority, count in priority_counts.items():
                    if count > 0:
                        # Convert priority name to lowercase for CSS class
                        class_name = priority.lower().replace("'", "").replace(" ", "-")
                        
                        # Add emoji based on priority
                        emoji = ""
                        if priority == "Highest":
                            emoji = "🔴"
                        elif priority == "High":
                            emoji = "🟠"
                        elif priority == "Medium":
                            emoji = "🟡"
                        elif priority == "Low":
                            emoji = "🔵"
                        elif priority == "Lowest":
                            emoji = "🔷"
                        elif priority == "Can't Automate":
                            emoji = "⚪"
                        
                        report_text += f"<span class='stat-pill {class_name}'>{emoji} {priority}: {count} tests</span>\n"
                
                report_text += "</div>\n</div>\n\n"
        
        return report_text

    @staticmethod
    def export_report_to_html(tests, priority_tiers, model, filename):
        """
        Export tests as a modern HTML report
        
        Args:
            tests (list): List of test dictionaries
            priority_tiers (dict): Dictionary with priority tier tests
            model: The prioritization model
            filename (str): Path to the output HTML file
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            from datetime import datetime
            
            # Count tests by priority
            priority_counts = {
                "Highest": len(priority_tiers["highest"]),
                "High": len(priority_tiers["high"]),
                "Medium": len(priority_tiers["medium"]),
                "Low": len(priority_tiers["low"]),
                "Lowest": len(priority_tiers["lowest"]),
                "Can't Automate": len(priority_tiers.get("cant_automate", []))
            }
            
            # Get thresholds
            highest_threshold = priority_tiers["highest_threshold"]
            high_threshold = priority_tiers["high_threshold"]
            medium_threshold = priority_tiers["medium_threshold"]
            low_threshold = priority_tiers["low_threshold"]
            
            # Start HTML content
            html = """<!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Test Automation Priority Report</title>
        <style>
            :root {
                --highest-color: #d32f2f;
                --high-color: #f57c00;
                --medium-color: #fbc02d;
                --low-color: #29b6f6;
                --lowest-color: #4fc3f7;
                --cant-automate-color: #9e9e9e;
                --bg-color: #f5f5f5;
                --card-bg: #ffffff;
                --text-color: #333333;
                --secondary-text: #666666;
                --border-color: #e0e0e0;
                --header-bg: #37474f;
                --header-text: #ffffff;
            }
            
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }
            
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: var(--text-color);
                background-color: var(--bg-color);
                padding: 0;
                margin: 0;
            }
            
            .container {
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
            }
            
            header {
                background-color: var(--header-bg);
                color: var(--header-text);
                padding: 20px 0;
                margin-bottom: 30px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            }
            
            header .container {
                display: flex;
                justify-content: space-between;
                align-items: center;
            }
            
            .header-content h1 {
                font-size: 24px;
                margin-bottom: 5px;
            }
            
            .timestamp {
                font-size: 14px;
                opacity: 0.8;
            }
            
            .summary-cards {
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(250px, 1fr));
                gap: 20px;
                margin-bottom: 30px;
            }
            
            .summary-card {
                background-color: var(--card-bg);
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 3px 10px rgba(0,0,0,0.1);
                display: flex;
                flex-direction: column;
            }
            
            .card-header {
                padding: 15px;
                font-weight: bold;
                color: white;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }
            
            .highest .card-header { background-color: var(--highest-color); }
            .high .card-header { background-color: var(--high-color); }
            .medium .card-header { background-color: var(--medium-color); }
            .low .card-header { background-color: var(--low-color); }
            .lowest .card-header { background-color: var(--lowest-color); }
            .cant-automate .card-header { background-color: var(--cant-automate-color); }
            
            .card-count {
                font-size: 24px;
                font-weight: bold;
            }
            
            .card-body {
                padding: 15px;
                font-size: 14px;
                color: var(--secondary-text);
                flex-grow: 1;
                display: flex;
                flex-direction: column;
                justify-content: space-between;
            }
            
            .card-threshold {
                font-size: 13px;
                margin-top: 10px;
            }
            
            section {
                margin-bottom: 40px;
            }
            
            .section-header {
                border-bottom: 2px solid var(--border-color);
                padding-bottom: 10px;
                margin-bottom: 20px;
                display: flex;
                align-items: center;
            }
            
            .section-icon {
                width: 24px;
                height: 24px;
                border-radius: 50%;
                margin-right: 10px;
                display: inline-block;
            }
            
            .section-highest .section-icon { background-color: var(--highest-color); }
            .section-high .section-icon { background-color: var(--high-color); }
            .section-medium .section-icon { background-color: var(--medium-color); }
            .section-low .section-icon { background-color: var(--low-color); }
            .section-lowest .section-icon { background-color: var(--lowest-color); }
            .section-cant-automate .section-icon { background-color: var(--cant-automate-color); }
            
            .section-description {
                font-size: 14px;
                color: var(--secondary-text);
                margin-bottom: 15px;
                font-style: italic;
            }
            
            .test-cards {
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(350px, 1fr));
                gap: 20px;
            }
            
            .test-card {
                background-color: var(--card-bg);
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 3px 10px rgba(0,0,0,0.1);
                transition: transform 0.2s ease-in-out;
            }
            
            .test-card:hover {
                transform: translateY(-5px);
            }
            
            .test-card-header {
                padding: 15px;
                border-bottom: 1px solid var(--border-color);
                position: relative;
            }
            
            .test-card-header h3 {
                margin-right: 50px;
                font-size: 16px;
            }
            
            .score-badge {
                position: absolute;
                top: 15px;
                right: 15px;
                width: 40px;
                height: 40px;
                border-radius: 50%;
                display: flex;
                align-items: center;
                justify-content: center;
                color: white;
                font-weight: bold;
            }
            
            .priority-highest .score-badge { background-color: var(--highest-color); }
            .priority-high .score-badge { background-color: var(--high-color); }
            .priority-medium .score-badge { background-color: var(--medium-color); }
            .priority-low .score-badge { background-color: var(--low-color); }
            .priority-lowest .score-badge { background-color: var(--lowest-color); }
            .priority-cant-automate .score-badge { background-color: var(--cant-automate-color); }
            
            .test-meta {
                display: flex;
                gap: 10px;
                margin-top: 5px;
                font-size: 13px;
            }
            
            .meta-item {
                background-color: var(--bg-color);
                border-radius: 4px;
                padding: 3px 8px;
            }
            
            .test-card-body {
                padding: 15px;
            }
            
            .test-description {
                margin-bottom: 15px;
                font-size: 14px;
                color: var(--secondary-text);
                border-left: 3px solid var(--border-color);
                padding-left: 10px;
            }
            
            .factor-list {
                list-style-type: none;
            }
            
            .factor-item {
                margin-bottom: 5px;
                font-size: 13px;
                display: flex;
                align-items: center;
            }
            
            .factor-score {
                display: inline-block;
                width: 20px;
                height: 20px;
                border-radius: 50%;
                background-color: var(--bg-color);
                color: var(--text-color);
                text-align: center;
                line-height: 20px;
                font-weight: bold;
                margin-right: 10px;
                font-size: 12px;
            }
            
            .factor-name {
                font-weight: bold;
                margin-right: 5px;
            }
            
            .factor-description {
                color: var(--secondary-text);
            }
            
            footer {
                text-align: center;
                padding: 20px;
                background-color: var(--header-bg);
                color: var(--header-text);
                font-size: 12px;
                margin-top: 50px;
            }
            
            .no-tests {
                text-align: center;
                padding: 50px;
                color: var(--secondary-text);
                font-style: italic;
            }
            
            @media (max-width: 768px) {
                .summary-cards {
                    grid-template-columns: repeat(auto-fill, minmax(150px, 1fr));
                }
                
                .test-cards {
                    grid-template-columns: 1fr;
                }
                
                header .container {
                    flex-direction: column;
                    text-align: center;
                }
                
                .header-meta {
                    margin-top: 10px;
                }
            }
        </style>
    </head>
    <body>
        <header>
            <div class="container">
                <div class="header-content">
                    <h1>Test Automation Priority Report</h1>
                    <div class="timestamp">Generated: """ + datetime.now().strftime('%Y-%m-%d %H:%M') + """</div>
                </div>
                <div class="header-meta">
                    <div>Total Tests: """ + str(len(tests)) + """</div>
                </div>
            </div>
        </header>
        
        <div class="container">
            <!-- Summary Cards -->
            <div class="summary-cards">
                <div class="summary-card highest">
                    <div class="card-header">
                        <span>Highest Priority</span>
                        <span class="card-count">""" + str(priority_counts["Highest"]) + """</span>
                    </div>
                    <div class="card-body">
                        <div>Recommended for immediate automation</div>
                        <div class="card-threshold">Score: ≥ """ + str(highest_threshold) + """</div>
                    </div>
                </div>
                
                <div class="summary-card high">
                    <div class="card-header">
                        <span>High Priority</span>
                        <span class="card-count">""" + str(priority_counts["High"]) + """</span>
                    </div>
                    <div class="card-body">
                        <div>Recommended for second phase</div>
                        <div class="card-threshold">Score: """ + str(high_threshold) + """ - """ + str(highest_threshold - 0.1) + """</div>
                    </div>
                </div>
                
                <div class="summary-card medium">
                    <div class="card-header">
                        <span>Medium Priority</span>
                        <span class="card-count">""" + str(priority_counts["Medium"]) + """</span>
                    </div>
                    <div class="card-body">
                        <div>Recommended for third phase</div>
                        <div class="card-threshold">Score: """ + str(medium_threshold) + """ - """ + str(high_threshold - 0.1) + """</div>
                    </div>
                </div>
                
                <div class="summary-card low">
                    <div class="card-header">
                        <span>Low Priority</span>
                        <span class="card-count">""" + str(priority_counts["Low"]) + """</span>
                    </div>
                    <div class="card-body">
                        <div>Consider for later phases</div>
                        <div class="card-threshold">Score: """ + str(low_threshold) + """ - """ + str(medium_threshold - 0.1) + """</div>
                    </div>
                </div>
                
                <div class="summary-card lowest">
                    <div class="card-header">
                        <span>Lowest Priority</span>
                        <span class="card-count">""" + str(priority_counts["Lowest"]) + """</span>
                    </div>
                    <div class="card-body">
                        <div>Not recommended for automation</div>
                        <div class="card-threshold">Score: < """ + str(low_threshold) + """</div>
                    </div>
                </div>
                
                <div class="summary-card cant-automate">
                    <div class="card-header">
                        <span>Can't Automate</span>
                        <span class="card-count">""" + str(priority_counts["Can't Automate"]) + """</span>
                    </div>
                    <div class="card-body">
                        <div>Not possible to automate</div>
                        <div class="card-threshold">Manual testing only</div>
                    </div>
                </div>
            </div>
    """
            
            # Add test sections by priority tier
            priority_sections = [
                ("Highest Priority Tests", priority_tiers["highest"], "highest", f"≥ {highest_threshold}", "Recommended for immediate automation"),
                ("High Priority Tests", priority_tiers["high"], "high", f"{high_threshold} - {highest_threshold - 0.1}", "Recommended for second phase automation"),
                ("Medium Priority Tests", priority_tiers["medium"], "medium", f"{medium_threshold} - {high_threshold - 0.1}", "Recommended for third phase automation"),
                ("Low Priority Tests", priority_tiers["low"], "low", f"{low_threshold} - {medium_threshold - 0.1}", "Consider for later phases or keep as manual tests"),
                ("Lowest Priority Tests", priority_tiers["lowest"], "lowest", f"< {low_threshold}", "Not recommended for automation")
            ]
            
            for title, tests_list, css_class, score_range, description in priority_sections:
                html += f"""
            <!-- {title} Section -->
            <section>
                <div class="section-header section-{css_class}">
                    <span class="section-icon"></span>
                    <h2>{title}</h2>
                </div>
                <div class="section-description">{description}. Score: {score_range}</div>
                
                <div class="test-cards">
    """
                
                # Add test cards
                if tests_list:
                    for i, test in enumerate(tests_list):
                        html += FileOperations._generate_test_card_html(test, i+1, css_class, model)
                else:
                    html += '<div class="no-tests">No tests in this category</div>'
                
                html += """
                </div>
            </section>
    """
            
            # Add "Can't Automate" section if there are tests that can't be automated
            if "cant_automate" in priority_tiers and priority_tiers["cant_automate"]:
                html += """
            <!-- Can't Automate Section -->
            <section>
                <div class="section-header section-cant-automate">
                    <span class="section-icon"></span>
                    <h2>Tests That Can't Be Automated</h2>
                </div>
                <div class="section-description">These tests have been identified as not possible to automate</div>
                
                <div class="test-cards">
    """
                
                for i, test in enumerate(priority_tiers["cant_automate"]):
                    html += FileOperations._generate_test_card_html(test, i+1, "cant-automate", model)
                
                html += """
                </div>
            </section>
    """
            
            # Close HTML
            html += """
        </div>
        
        <footer>
            <div>Test Automation Priority Report</div>
            <div>Generated using the Test Prioritization Tool</div>
        </footer>
    </body>
    </html>"""
            
            # Write to file
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(html)
            
            return True
        
        except Exception as e:
            print(f"HTML export error: {e}")
            return False

    @staticmethod
    def _generate_test_card_html(test, index, priority_class, model):
        """
        Generate HTML for a test card
        
        Args:
            test (dict): Test dictionary
            index (int): Index/rank of the test
            priority_class (str): CSS class for priority styling
            model: The prioritization model for factor descriptions
        
        Returns:
            str: HTML for the test card
        """
        try:
            # Start card
            card = f"""
                    <div class="test-card priority-{priority_class}">
                        <div class="test-card-header">
                            <h3>{index}. {test['name']}</h3>
                            <span class="score-badge">{test['total_score']}</span>
                            <div class="test-meta">
                                <span class="meta-item ticket">{test.get('ticket_id', 'N/A')}</span>
            """
            
            # Add section if available
            if test.get("section"):
                card += f'<span class="meta-item section">{test["section"]}</span>'
            
            card += """
                            </div>
                        </div>
                        <div class="test-card-body">
            """
            
            # Add description if available
            if test.get('description'):
                card += f'<div class="test-description">{test["description"]}</div>'
            
            # Start factor list
            card += '<ul class="factor-list">'
            
            # Add factors
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                # Show the "Can it be automated?" factor first if in "Can't Automate" category
                if test['priority'] == "Can't Automate" and "can_be_automated" in test['scores']:
                    factor_key = "can_be_automated"
                    factor_name = model.factors[factor_key]["name"]
                    score = test['scores'][factor_key]
                    score_description = model.score_options[factor_key][score]
                    
                    card += f"""
                            <li class="factor-item">
                                <span class="factor-score">{score}</span>
                                <span class="factor-name">{factor_name}:</span>
                                <span class="factor-description">{score_description}</span>
                            </li>
                    """
                
                # Add other factors
                for factor_key, score in test['scores'].items():
                    # Skip the can_be_automated factor if already shown or if test can't be automated
                    if factor_key == "can_be_automated" and (test['priority'] == "Can't Automate"):
                        continue
                        
                    if factor_key in model.factors and score in model.score_options.get(factor_key, {}):
                        factor_name = model.factors[factor_key]["name"]
                        score_description = model.score_options[factor_key][score]
                        
                        card += f"""
                            <li class="factor-item">
                                <span class="factor-score">{score}</span>
                                <span class="factor-name">{factor_name}:</span>
                                <span class="factor-description">{score_description}</span>
                            </li>
                        """
            
            # Close factor list and card
            card += """
                        </ul>
                    </div>
                </div>
            """
            
            return card
        except Exception as e:
            print(f"Error generating test card: {e}")
            return f"<div>Error generating card for {test.get('name', 'Unknown Test')}</div>"
    
    @staticmethod
    def generate_scoring_guide_text(factors, score_options):
        """
        Generate text for the scoring guide
        
        Args:
            factors (dict): Dictionary of factors and weights
            score_options (dict): Dictionary of scoring options
            
        Returns:
            str: Formatted guide text
        """
        guide = "TEST AUTOMATION PRIORITIZATION SCORING GUIDE\n"
        guide += "=" * 44 + "\n\n"
        
        guide += "This application uses the following weighted factors to calculate\n"
        guide += "which manual tests should be prioritized for automation:\n\n"
        
        for factor, details in factors.items():
            factor_name = details["name"]
            weight = details["weight"]
            
            guide += f"{factor_name} (Weight: {weight})\n"
            guide += "-" * 50 + "\n"
            
            for score, description in score_options[factor].items():
                guide += f"  {score} - {description}\n"
            
            guide += "\n"
        
        guide += "How scores are calculated:\n"
        guide += "-" * 50 + "\n"
        guide += "1. Each factor score is multiplied by its weight\n"
        guide += "2. These weighted scores are summed to get a raw score\n"
        guide += "3. The raw score is converted to a 100-point scale\n\n"
        guide += "Formula: Final Score = (Raw Score / Max Possible Raw Score) × 100\n\n"
        
        # Skip the "can_be_automated" factor (which has weight 0) when calculating max possible score
        max_possible_raw = sum(5 * details["weight"] for factor, details in factors.items() 
                            if factor != "can_be_automated")
        guide += f"Maximum possible raw score: {max_possible_raw}\n"
        guide += "Maximum possible final score: 100\n"
        guide += f"Minimum possible raw score: {sum(1 * details['weight'] for factor, details in factors.items() if factor != 'can_be_automated')}\n"
        guide += "Minimum possible final score: 20\n\n"
        
        # Special note for "Can it be automated?" factor
        if "can_be_automated" in factors:
            guide += "Special Case - Tests that cannot be automated:\n"
            guide += "-" * 50 + "\n"
            guide += "If a test is marked as 'Cannot be automated' (selecting 'No' for the\n"
            guide += "'Can it be automated?' factor), it will automatically receive:\n"
            guide += "  - A score of 0\n"
            guide += "  - Priority category of 'Can't Automate'\n"
            guide += "These tests are excluded from normal prioritization and shown separately.\n\n"
        
        return guide
    
    @staticmethod
    def export_report_to_file(filename, report_text):
        """
        Export a report to a text file
        
        Args:
            filename (str): Path to the text file
            report_text (str): The report text to export
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            with open(filename, "w") as f:
                f.write(report_text)
            return True
        except Exception as e:
            print(f"Export report error: {str(e)}")
            return False
        
    @staticmethod
    def export_report_to_html(tests, priority_tiers, model, filename):
        """
        Export tests as a modern HTML report
        
        Args:
            tests (list): List of test dictionaries
            priority_tiers (dict): Dictionary with priority tier tests
            model: The prioritization model
            filename (str): Path to the output HTML file
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            from datetime import datetime
            
            # Count tests by priority
            priority_counts = {
                "Highest": len(priority_tiers["highest"]),
                "High": len(priority_tiers["high"]),
                "Medium": len(priority_tiers["medium"]),
                "Low": len(priority_tiers["low"]),
                "Lowest": len(priority_tiers["lowest"]),
                "Can't Automate": len(priority_tiers.get("cant_automate", []))
            }
            
            # Get thresholds
            highest_threshold = priority_tiers["highest_threshold"]
            high_threshold = priority_tiers["high_threshold"]
            medium_threshold = priority_tiers["medium_threshold"]
            low_threshold = priority_tiers["low_threshold"]
            
            # Start HTML content
            html = f"""<!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Test Automation Priority Report</title>
        <style>
            :root {{
                --highest-color: #d32f2f;
                --high-color: #f57c00;
                --medium-color: #fbc02d;
                --low-color: #29b6f6;
                --lowest-color: #4fc3f7;
                --cant-automate-color: #9e9e9e;
                --bg-color: #f5f5f5;
                --card-bg: #ffffff;
                --text-color: #333333;
                --secondary-text: #666666;
                --border-color: #e0e0e0;
                --header-bg: #37474f;
                --header-text: #ffffff;
            }}
            
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                color: var(--text-color);
                background-color: var(--bg-color);
                padding: 0;
                margin: 0;
            }}
            
            .container {{
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
            }}
            
            header {{
                background-color: var(--header-bg);
                color: var(--header-text);
                padding: 20px 0;
                margin-bottom: 30px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            }}
            
            header .container {{
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            
            .header-content h1 {{
                font-size: 24px;
                margin-bottom: 5px;
            }}
            
            .timestamp {{
                font-size: 14px;
                opacity: 0.8;
            }}
            
            .summary-cards {{
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(250px, 1fr));
                gap: 20px;
                margin-bottom: 30px;
            }}
            
            .summary-card {{
                background-color: var(--card-bg);
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 3px 10px rgba(0,0,0,0.1);
                display: flex;
                flex-direction: column;
            }}
            
            .card-header {{
                padding: 15px;
                font-weight: bold;
                color: white;
                display: flex;
                justify-content: space-between;
                align-items: center;
            }}
            
            .highest .card-header {{ background-color: var(--highest-color); }}
            .high .card-header {{ background-color: var(--high-color); }}
            .medium .card-header {{ background-color: var(--medium-color); }}
            .low .card-header {{ background-color: var(--low-color); }}
            .lowest .card-header {{ background-color: var(--lowest-color); }}
            .cant-automate .card-header {{ background-color: var(--cant-automate-color); }}
            
            .card-count {{
                font-size: 24px;
                font-weight: bold;
            }}
            
            .card-body {{
                padding: 15px;
                font-size: 14px;
                color: var(--secondary-text);
                flex-grow: 1;
                display: flex;
                flex-direction: column;
                justify-content: space-between;
            }}
            
            .card-threshold {{
                font-size: 13px;
                margin-top: 10px;
            }}
            
            section {{
                margin-bottom: 40px;
            }}
            
            .section-header {{
                border-bottom: 2px solid var(--border-color);
                padding-bottom: 10px;
                margin-bottom: 20px;
                display: flex;
                align-items: center;
            }}
            
            .section-icon {{
                width: 24px;
                height: 24px;
                border-radius: 50%;
                margin-right: 10px;
                display: inline-block;
            }}
            
            .section-highest .section-icon {{ background-color: var(--highest-color); }}
            .section-high .section-icon {{ background-color: var(--high-color); }}
            .section-medium .section-icon {{ background-color: var(--medium-color); }}
            .section-low .section-icon {{ background-color: var(--low-color); }}
            .section-lowest .section-icon {{ background-color: var(--lowest-color); }}
            .section-cant-automate .section-icon {{ background-color: var(--cant-automate-color); }}
            
            .section-description {{
                font-size: 14px;
                color: var(--secondary-text);
                margin-bottom: 15px;
                font-style: italic;
            }}
            
            .test-cards {{
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(350px, 1fr));
                gap: 20px;
            }}
            
            .test-card {{
                background-color: var(--card-bg);
                border-radius: 8px;
                overflow: hidden;
                box-shadow: 0 3px 10px rgba(0,0,0,0.1);
                transition: transform 0.2s ease-in-out;
            }}
            
            .test-card:hover {{
                transform: translateY(-5px);
            }}
            
            .test-card-header {{
                padding: 15px;
                border-bottom: 1px solid var(--border-color);
                position: relative;
            }}
            
            .test-card-header h3 {{
                margin-right: 50px;
                font-size: 16px;
            }}
            
            .score-badge {{
                position: absolute;
                top: 15px;
                right: 15px;
                width: 40px;
                height: 40px;
                border-radius: 50%;
                display: flex;
                align-items: center;
                justify-content: center;
                color: white;
                font-weight: bold;
            }}
            
            .priority-highest .score-badge {{ background-color: var(--highest-color); }}
            .priority-high .score-badge {{ background-color: var(--high-color); }}
            .priority-medium .score-badge {{ background-color: var(--medium-color); }}
            .priority-low .score-badge {{ background-color: var(--low-color); }}
            .priority-lowest .score-badge {{ background-color: var(--lowest-color); }}
            .priority-cant-automate .score-badge {{ background-color: var(--cant-automate-color); }}
            
            .test-meta {{
                display: flex;
                gap: 10px;
                margin-top: 5px;
                font-size: 13px;
            }}
            
            .meta-item {{
                background-color: var(--bg-color);
                border-radius: 4px;
                padding: 3px 8px;
            }}
            
            .test-card-body {{
                padding: 15px;
            }}
            
            .test-description {{
                margin-bottom: 15px;
                font-size: 14px;
                color: var(--secondary-text);
                border-left: 3px solid var(--border-color);
                padding-left: 10px;
            }}
            
            .factor-list {{
                list-style-type: none;
            }}
            
            .factor-item {{
                margin-bottom: 5px;
                font-size: 13px;
                display: flex;
                align-items: center;
            }}
            
            .factor-score {{
                display: inline-block;
                width: 20px;
                height: 20px;
                border-radius: 50%;
                background-color: var(--bg-color);
                color: var(--text-color);
                text-align: center;
                line-height: 20px;
                font-weight: bold;
                margin-right: 10px;
                font-size: 12px;
            }}
            
            .factor-name {{
                font-weight: bold;
                margin-right: 5px;
            }}
            
            .factor-description {{
                color: var(--secondary-text);
            }}
            
            footer {{
                text-align: center;
                padding: 20px;
                background-color: var(--header-bg);
                color: var(--header-text);
                font-size: 12px;
                margin-top: 50px;
            }}
            
            .no-tests {{
                text-align: center;
                padding: 50px;
                color: var(--secondary-text);
                font-style: italic;
            }}
            
            @media (max-width: 768px) {{
                .summary-cards {{
                    grid-template-columns: repeat(auto-fill, minmax(150px, 1fr));
                }}
                
                .test-cards {{
                    grid-template-columns: 1fr;
                }}
                
                header .container {{
                    flex-direction: column;
                    text-align: center;
                }}
                
                .header-meta {{
                    margin-top: 10px;
                }}
            }}
        </style>
    </head>
    <body>
        <header>
            <div class="container">
                <div class="header-content">
                    <h1>Test Automation Priority Report</h1>
                    <div class="timestamp">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
                </div>
                <div class="header-meta">
                    <div>Total Tests: {len(tests)}</div>
                </div>
            </div>
        </header>
        
        <div class="container">
            <!-- Summary Cards -->
            <div class="summary-cards">
                <div class="summary-card highest">
                    <div class="card-header">
                        <span>Highest Priority</span>
                        <span class="card-count">{priority_counts["Highest"]}</span>
                    </div>
                    <div class="card-body">
                        <div>Recommended for immediate automation</div>
                        <div class="card-threshold">Score: ≥ {highest_threshold}</div>
                    </div>
                </div>
                
                <div class="summary-card high">
                    <div class="card-header">
                        <span>High Priority</span>
                        <span class="card-count">{priority_counts["High"]}</span>
                    </div>
                    <div class="card-body">
                        <div>Recommended for second phase</div>
                        <div class="card-threshold">Score: {high_threshold} - {highest_threshold - 0.1}</div>
                    </div>
                </div>
                
                <div class="summary-card medium">
                    <div class="card-header">
                        <span>Medium Priority</span>
                        <span class="card-count">{priority_counts["Medium"]}</span>
                    </div>
                    <div class="card-body">
                        <div>Recommended for third phase</div>
                        <div class="card-threshold">Score: {medium_threshold} - {high_threshold - 0.1}</div>
                    </div>
                </div>
                
                <div class="summary-card low">
                    <div class="card-header">
                        <span>Low Priority</span>
                        <span class="card-count">{priority_counts["Low"]}</span>
                    </div>
                    <div class="card-body">
                        <div>Consider for later phases</div>
                        <div class="card-threshold">Score: {low_threshold} - {medium_threshold - 0.1}</div>
                    </div>
                </div>
                
                <div class="summary-card lowest">
                    <div class="card-header">
                        <span>Lowest Priority</span>
                        <span class="card-count">{priority_counts["Lowest"]}</span>
                    </div>
                    <div class="card-body">
                        <div>Not recommended for automation</div>
                        <div class="card-threshold">Score: < {low_threshold}</div>
                    </div>
                </div>
                
                <div class="summary-card cant-automate">
                    <div class="card-header">
                        <span>Can't Automate</span>
                        <span class="card-count">{priority_counts["Can't Automate"]}</span>
                    </div>
                    <div class="card-body">
                        <div>Not possible to automate</div>
                        <div class="card-threshold">Manual testing only</div>
                    </div>
                </div>
            </div>
    """
            
            # Add test sections by priority tier
            priority_sections = [
                ("Highest Priority Tests", priority_tiers["highest"], "highest", f"≥ {highest_threshold}", "Recommended for immediate automation"),
                ("High Priority Tests", priority_tiers["high"], "high", f"{high_threshold} - {highest_threshold - 0.1}", "Recommended for second phase automation"),
                ("Medium Priority Tests", priority_tiers["medium"], "medium", f"{medium_threshold} - {high_threshold - 0.1}", "Recommended for third phase automation"),
                ("Low Priority Tests", priority_tiers["low"], "low", f"{low_threshold} - {medium_threshold - 0.1}", "Consider for later phases or keep as manual tests"),
                ("Lowest Priority Tests", priority_tiers["lowest"], "lowest", f"< {low_threshold}", "Not recommended for automation")
            ]
            
            for title, tests_list, css_class, score_range, description in priority_sections:
                html += f"""
            <!-- {title} Section -->
            <section>
                <div class="section-header section-{css_class}">
                    <span class="section-icon"></span>
                    <h2>{title}</h2>
                </div>
                <div class="section-description">{description}. Score: {score_range}</div>
                
                <div class="test-cards">
    """
                
                # Add test cards
                if tests_list:
                    for i, test in enumerate(tests_list):
                        html += FileOperations._generate_test_card_html(test, i+1, css_class, model)
                else:
                    html += '<div class="no-tests">No tests in this category</div>'
                
                html += """
                </div>
            </section>
    """
            
            # Add "Can't Automate" section if there are tests that can't be automated
            if "cant_automate" in priority_tiers and priority_tiers["cant_automate"]:
                html += """
            <!-- Can't Automate Section -->
            <section>
                <div class="section-header section-cant-automate">
                    <span class="section-icon"></span>
                    <h2>Tests That Can't Be Automated</h2>
                </div>
                <div class="section-description">These tests have been identified as not possible to automate</div>
                
                <div class="test-cards">
    """
                
                for i, test in enumerate(priority_tiers["cant_automate"]):
                    html += FileOperations._generate_test_card_html(test, i+1, "cant-automate", model)
                
                html += """
                </div>
            </section>
    """
            
            # Close HTML
            html += """
        </div>
        
        <footer>
            <div>Test Automation Priority Report</div>
            <div>Generated using the Test Prioritization Tool</div>
        </footer>
    </body>
    </html>"""
            
            # Write to file
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(html)
            
            return True
        
        except Exception as e:
            print(f"HTML export error: {e}")
            return False

    @staticmethod
    def _generate_test_card_html(test, index, priority_class, model):
        """
        Generate HTML for a test card
        
        Args:
            test (dict): Test dictionary
            index (int): Index/rank of the test
            priority_class (str): CSS class for priority styling
            model: The prioritization model for factor descriptions
        
        Returns:
            str: HTML for the test card
        """
        try:
            # Start card
            card = f"""
                    <div class="test-card priority-{priority_class}">
                        <div class="test-card-header">
                            <h3>{index}. {test['name']}</h3>
                            <span class="score-badge">{test['total_score']}</span>
                            <div class="test-meta">
                                <span class="meta-item ticket">{test.get('ticket_id', 'N/A')}</span>
            """
            
            # Add section if available
            if test.get("section"):
                card += f'<span class="meta-item section">{test["section"]}</span>'
            
            card += """
                            </div>
                        </div>
                        <div class="test-card-body">
            """
            
            # Add description if available
            if test.get('description'):
                card += f'<div class="test-description">{test["description"]}</div>'
            
            # Start factor list
            card += '<ul class="factor-list">'
            
            # Add factors
            if model and hasattr(model, 'factors') and hasattr(model, 'score_options'):
                # Show the "Can it be automated?" factor first if in "Can't Automate" category
                if test['priority'] == "Can't Automate" and "can_be_automated" in test['scores']:
                    factor_key = "can_be_automated"
                    factor_name = model.factors[factor_key]["name"]
                    score = test['scores'][factor_key]
                    score_description = model.score_options[factor_key][score]
                    
                    card += f"""
                            <li class="factor-item">
                                <span class="factor-score">{score}</span>
                                <span class="factor-name">{factor_name}:</span>
                                <span class="factor-description">{score_description}</span>
                            </li>
                    """
                
                # Add other factors
                for factor_key, score in test['scores'].items():
                    # Skip the can_be_automated factor if already shown or if test can't be automated
                    if factor_key == "can_be_automated" and (test['priority'] == "Can't Automate"):
                        continue
                        
                    if factor_key in model.factors and score in model.score_options.get(factor_key, {}):
                        factor_name = model.factors[factor_key]["name"]
                        score_description = model.score_options[factor_key][score]
                        
                        card += f"""
                            <li class="factor-item">
                                <span class="factor-score">{score}</span>
                                <span class="factor-name">{factor_name}:</span>
                                <span class="factor-description">{score_description}</span>
                            </li>
                        """
            
            # Close factor list and card
            card += """
                        </ul>
                    </div>
                </div>
            """
            
            return card
        except Exception as e:
            print(f"Error generating test card: {e}")
            return f"<div>Error generating card for {test.get('name', 'Unknown Test')}</div>"

    @staticmethod
    def export_report_to_docx(report_text, filename):
        """
        Export markdown report to Word (.docx) file
        
        Args:
            report_text (str): Markdown formatted report text
            filename (str): Path to the output .docx file
            
        Returns:
            tuple: (success, error_message)
                success (bool): True if successful, False otherwise
                error_message (str): Error message if unsuccessful, None otherwise
        """
        try:
            # Try to import python-docx module
            try:
                import docx
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                has_docx = True
            except ImportError:
                has_docx = False
                
            if not has_docx:
                return False, "python-docx module not installed. Install it with 'pip install python-docx'"
            
            # Create a new document
            doc = docx.Document()
            
            # Define styles for different heading levels and text
            styles = {
                'title': {'size': 16, 'bold': True, 'color': RGBColor(44, 62, 80)},
                'heading1': {'size': 14, 'bold': True, 'color': RGBColor(52, 152, 219)},
                'heading2': {'size': 12, 'bold': True, 'color': RGBColor(41, 128, 185)},
                'heading3': {'size': 11, 'bold': True, 'color': RGBColor(36, 113, 163)},
                'normal': {'size': 10, 'bold': False, 'color': RGBColor(0, 0, 0)},
                'emphasis': {'size': 10, 'bold': False, 'italic': True, 'color': RGBColor(0, 0, 0)},
                'strong': {'size': 10, 'bold': True, 'color': RGBColor(0, 0, 0)}
            }
            
            # Process markdown line by line
            lines = report_text.split('\n')
            in_list = False
            list_items = []
            
            for line in lines:
                # Skip empty lines
                if not line.strip():
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                    doc.add_paragraph()
                    continue
                
                # Handle headings
                if line.startswith('# '):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:])
                    font = run.font
                    font.size = Pt(styles['title']['size'])
                    font.bold = styles['title']['bold']
                    font.color.rgb = styles['title']['color']
                    
                elif line.startswith('## '):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    font = run.font
                    font.size = Pt(styles['heading1']['size'])
                    font.bold = styles['heading1']['bold']
                    font.color.rgb = styles['heading1']['color']
                    
                elif line.startswith('### '):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    font = run.font
                    font.size = Pt(styles['heading2']['size'])
                    font.bold = styles['heading2']['bold']
                    font.color.rgb = styles['heading2']['color']
                    
                # Handle horizontal rule
                elif line.startswith('---'):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph('_' * 50)
                    
                # Handle list items
                elif line.startswith('* '):
                    in_list = True
                    list_items.append(line[2:])
                    
                # Handle normal paragraphs
                else:
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                    
                    # Handle bold and italic formatting
                    p = doc.add_paragraph()
                    
                    # Replace emojis with their names
                    line = (line.replace('🔴', '[HIGHEST] ')
                            .replace('🟠', '[HIGH] ')
                            .replace('🟡', '[MEDIUM] ')
                            .replace('🔵', '[LOW] ')
                            .replace('🔷', '[LOWEST] ')
                            .replace('⚪', '[NOT AUTOMATABLE] '))
                    
                    # Process bold sections
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Bold text (odd indices)
                            run = p.add_run(part)
                            font = run.font
                            font.bold = True
                        else:  # Normal text (even indices)
                            # Process italic sections within normal text
                            italic_parts = part.split('*')
                            for j, italic_part in enumerate(italic_parts):
                                if j % 2 == 1:  # Italic text (odd indices)
                                    run = p.add_run(italic_part)
                                    font = run.font
                                    font.italic = True
                                else:  # Regular text (even indices)
                                    if italic_part:
                                        run = p.add_run(italic_part)
            
            # End any open list
            if in_list:
                for item in list_items:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item)
            
            # Save the document
            doc.save(filename)
            return True, None
            
        except Exception as e:
            print(f"Word export error: {str(e)}")
            return False, str(e)