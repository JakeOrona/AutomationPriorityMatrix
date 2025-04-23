"""
docx_report.py - Word document report generator for test prioritization
"""
from .base_report import BaseReportGenerator

class DocxReportGenerator(BaseReportGenerator):
    """
    Handles Word document (.docx) report generation for the test prioritization application
    """

    @staticmethod
    def export_report(report_text, filename):
        """
        Export markdown report to Word (.docx) file
        
        Args:
            report_text (str): Markdown formatted report text
            filename (str): Path to the output .docx file
            
        Returns:
            tuple: (success, error_message)
                success (bool): True if successful, False otherwise
                error_message (str): Error message if unsuccessful, None otherwise
        """
        try:
            # Try to import python-docx module
            try:
                import docx
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                has_docx = True
            except ImportError:
                has_docx = False
                
            if not has_docx:
                return False, "python-docx module not installed. Install it with 'pip install python-docx'"
            
            # Create a new document
            doc = docx.Document()
            
            # Define styles for different heading levels and text
            styles = {
                'title': {'size': 16, 'bold': True, 'color': RGBColor(44, 62, 80)},
                'heading1': {'size': 14, 'bold': True, 'color': RGBColor(52, 152, 219)},
                'heading2': {'size': 12, 'bold': True, 'color': RGBColor(41, 128, 185)},
                'heading3': {'size': 11, 'bold': True, 'color': RGBColor(36, 113, 163)},
                'normal': {'size': 10, 'bold': False, 'color': RGBColor(0, 0, 0)},
                'emphasis': {'size': 10, 'bold': False, 'italic': True, 'color': RGBColor(0, 0, 0)},
                'strong': {'size': 10, 'bold': True, 'color': RGBColor(0, 0, 0)}
            }
            
            # Process markdown line by line
            lines = report_text.split('\n')
            in_list = False
            list_items = []
            
            for line in lines:
                # Skip empty lines
                if not line.strip():
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                    doc.add_paragraph()
                    continue
                
                # Handle headings
                if line.startswith('# '):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:])
                    font = run.font
                    font.size = Pt(styles['title']['size'])
                    font.bold = styles['title']['bold']
                    font.color.rgb = styles['title']['color']
                    
                elif line.startswith('## '):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    font = run.font
                    font.size = Pt(styles['heading1']['size'])
                    font.bold = styles['heading1']['bold']
                    font.color.rgb = styles['heading1']['color']
                    
                elif line.startswith('### '):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    font = run.font
                    font.size = Pt(styles['heading2']['size'])
                    font.bold = styles['heading2']['bold']
                    font.color.rgb = styles['heading2']['color']
                    
                # Handle horizontal rule
                elif line.startswith('---'):
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                        
                    p = doc.add_paragraph('_' * 50)
                    
                # Handle list items
                elif line.startswith('* '):
                    in_list = True
                    list_items.append(line[2:])
                    
                # Handle normal paragraphs
                else:
                    if in_list:
                        # End current list
                        for item in list_items:
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(item)
                        list_items = []
                        in_list = False
                    
                    # Handle bold and italic formatting
                    p = doc.add_paragraph()
                    
                    # Replace emojis with their names
                    line = (line.replace('🔴', '[HIGHEST] ')
                            .replace('🟠', '[HIGH] ')
                            .replace('🟡', '[MEDIUM] ')
                            .replace('🔵', '[LOW] ')
                            .replace('🔷', '[LOWEST] ')
                            .replace('⚪', '[NOT AUTOMATABLE] '))
                    
                    # Process bold sections
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # Bold text (odd indices)
                            run = p.add_run(part)
                            font = run.font
                            font.bold = True
                        else:  # Normal text (even indices)
                            # Process italic sections within normal text
                            italic_parts = part.split('*')
                            for j, italic_part in enumerate(italic_parts):
                                if j % 2 == 1:  # Italic text (odd indices)
                                    run = p.add_run(italic_part)
                                    font = run.font
                                    font.italic = True
                                else:  # Regular text (even indices)
                                    if italic_part:
                                        run = p.add_run(italic_part)
            
            # End any open list
            if in_list:
                for item in list_items:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item)
            
            # Save the document
            doc.save(filename)
            return True, None
            
        except Exception as e:
            error_message = f"Word export error: {str(e)}"
            print(error_message)
            return False, error_message